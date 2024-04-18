from tkinter import filedialog
import pandas as pd
import docx as doc
import tkinter as tk
from tkinter import ttk
from docx.oxml.ns import qn
import os
import json
print('''
 ______   __  __     ______        ______     ______     ______     ______     ______   ______     ______
/\__  _\ /\ \_\ \   /\  ___\      /\  ___\   /\  == \   /\  ___\   /\  __ \   /\__  _\ /\  __ \   /\  == \
\/_/\ \/ \ \  __ \  \ \  __\      \ \ \____  \ \  __<   \ \  __\   \ \  __ \  \/_/\ \/ \ \ \/\ \  \ \  __<
   \ \_\  \ \_\ \_\  \ \_____\     \ \_____\  \ \_\ \_\  \ \_____\  \ \_\ \_\    \ \_\  \ \_____\  \ \_\ \_\
    \/_/   \/_/\/_/   \/_____/      \/_____/   \/_/ /_/   \/_____/   \/_/\/_/     \/_/   \/_____/   \/_/ /_/
                                                                                                            ''')
print('Waiting to build.........')
print('等待构建中......')
'''
提取用户可以手动输入的变量
1.读取的excel文件路径
2.读取的工作表
3.常用岗位
4.word文档保存路径
'''

EXCEL_PATH = ""
SHEET_NAME = "Sheet1"
POSISION_TITLE = ["指挥长", "指挥助理", "调度助理", "通信助理", "保障助理", "政宣助理", "行政值班"]
DOC_SAVE_PATH = ""
TOBE_SELECTED = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12", "13", "14", "15",
    "16", "17", "18", "19", "20", "21", "22", "23", "24", "25", "26", "27", "28", "29", "30", "31"]
DATAFRAME_SAVE =""
CREATE_FOLDER_PATH="C:\\ProgramData\\AutoCreateDuty"
CREATE_JSON_PATH = "C:\\ProgramData\\AutoCreateDuty\\TobeFiledItem.json"


class CreateTable:
     global EXCEL_PATH  #excel表存放的路径
     global SHEET_NAME  #读取excel表的工作表名称
     global POSISION_TITLE #常用岗位的名称
     global DOC_SAVE_PATH #word文档保存的路径
     global TOBE_SELECTED #用于选择日期的列表
     global DATAFRAME_SAVE #保存读取的excel表对象
     global CREATE_FOLDER_PATH
     global CREATE_JSON_PATH
     Template = '''
<日期>总队全勤指挥部值班人员
【带班领导】<带班领导>
【总指挥】<总指挥>
【指挥长】<指挥长><指挥长备勤>（备勤）
【指挥助理】<指挥助理><指挥助理备勤>（备勤）
【调度助理】<调度助理><调度助理备勤>（备勤）
【通信助理】<通信助理><通信助理备勤>（备勤）
【保障助理】<保障助理><保障助理备勤>（备勤）
【政宣助理】<政宣助理><政宣助理备勤>（备勤）
【行政值班】<行政值班><行政值班备勤>（备勤）
【值班专家】<值班专家>
【值班驾驶员】<值班驾驶员>
【通信分队】<通信分队>
总队灭火救援指挥专班值班人员：
<灭火救援指挥专班指挥长>、<灭火救援指挥专班指挥助理>、<灭火救援指挥专班信息助理>、<灭火救援指挥专班通信助理>
【协作区轮值指挥长】
鲁中协作区 <鲁中协作区>
鲁东协作区 <鲁东协作区>
鲁南协作区 <鲁南协作区>
【值班提示】
值班期间，请值班人员不要离开营区，备班人员不要离开辖区，保持通信畅通，接指挥中心调派通知后立即响应。各岗位值班人员换班须报值班领导、指挥长批准并向指挥中心备案。

总指挥：<总指挥>
（值班提示）报告首长：今日指挥长<指挥长>，指挥助理<指挥助理>，调度助理<调度助理>，通信助理<通信助理>，保障助理<保障助理>，政宣助理<政宣助理>，行政值班<行政值班>，值班专家<值班专家>，通信分队<通信分队>，值班人员全部在岗在位。（总队指挥中心）
'''

     def __init__(self):
        # 引入全局变量
        global DOC_SAVE_PATH
        # 获取到当前用户的桌面路径，修改全局变量为桌面路径+创建的文件夹
        DOC_SAVE_PATH = (os.path.join(os.path.join(
            os.environ['USERPROFILE']), 'Desktop'))+"\\值班短信自动生成_"
        # 创建UI
        self.initUI()

        

     def initUI(self):
          self.myTk = tk.Tk()
          self.myTk.title("值班信息自动生成 @Author:THE CREATOR    v1.1.0")
          self.myTk.geometry("460x350")
          # 禁止用户缩放窗体，第一个参数是水平缩放，第二个参数是垂直缩放
          # self.myTk.resizable(False, False)

          # 创建选择文件按钮
          self.ChoiceBtn = tk.Button(
              self.myTk, text="   选择文件   ", command=self.handleFilePath)
          # -self.ChoiceBtn.pack()
          self.ChoiceBtn.grid(row=1, column=0, columnspan=4)

          # 创建开始生成按钮
          self.CreateBtn = tk.Button(
              self.myTk, text="生成全月值班信息", command=self.WriteAllfiedDateToDoc)
          # -self.CreateBtn.pack()
          self.CreateBtn.grid(row=1, column=4, columnspan=2, pady=10)
          # 创建单日期输出的Label
          self.SingleDateLabel = tk.Label(self.myTk, text="选择指定日期进行输出：")
          self.SingleDateLabel.grid(
              row=0, column=0, columnspan=2, pady=20, padx=10)
          # 创建选择日期下拉选项
          self.SelectDateBox = ttk.Combobox(
              self.myTk, values=TOBE_SELECTED, width=7)
          self.SelectDateBox.current(0)
          self.SelectDateBox.grid(row=0, column=2, columnspan=2)

          # 创建开始单日期输出按钮
          self.OutalongDateDoc = tk.Button(
              self.myTk, text="指定日期输出", command=self.WriteSpecifiedDateToDoc)
          self.OutalongDateDoc.grid(row=0, column=4, columnspan=2, padx=8)

          # 创建值班专家Label
          self.SpecialistLabel = tk.Label(self.myTk, text="值班专家：")
          # -self.SpecialistLabel.pack()
          self.SpecialistLabel.grid(row=2, column=0, columnspan=2, sticky="E")
          # 创建值班专家输入框
          self.SpecialistInput = tk.Entry(self.myTk)
          # -self.SpecialistInput.pack()
          self.SpecialistInput.grid(row=2, column=2, columnspan=4, pady=4)

          # 创建值班驾驶员Label
          self.DriverLabel = tk.Label(self.myTk, text="值班驾驶员：")
          # -self.DriverLabel.pack()
          self.DriverLabel.grid(row=3, column=0, columnspan=2, sticky="E")
          # 创建值班驾驶员输入框
          self.DriverInput = tk.Entry(self.myTk)
          # -self.DriverInput.pack()
          self.DriverInput.grid(row=3, column=2, columnspan=4, pady=4)

          # 创建鲁中协作区Label
          self.LZLabel = tk.Label(self.myTk, text="鲁中协作区：")
          # -self.LZLabel.pack()
          self.LZLabel.grid(row=4, column=0, columnspan=2, sticky="E")
          # 创建鲁中协作区输入框
          self.LZInput = tk.Entry(self.myTk)
          # -self.LZInput.pack()
          self.LZInput.grid(row=4, column=2, columnspan=4, pady=4)

          # 创建鲁东协作区Label
          self.LDLabel = tk.Label(self.myTk, text="鲁东协作区：")
          # -self.LDLabel.pack()
          self.LDLabel.grid(row=5, column=0, columnspan=2, sticky="E")

          # 创建鲁东协作区输入框
          self.LDInput = tk.Entry(self.myTk)
          # -self.LDInput.pack()
          self.LDInput.grid(row=5, column=2, columnspan=4, pady=4)

          # 创建鲁南协作区Label
          self.LNLabel = tk.Label(self.myTk, text="鲁南协作区：")
          # -self.LNLabel.pack()
          self.LNLabel.grid(row=6, column=0, columnspan=2, sticky="E")
          # 创建鲁南协作区输入框
          self.LNInput = tk.Entry(self.myTk)
          # -self.LNInput.pack()
          self.LNInput.grid(row=6, column=2, columnspan=4, pady=4)

          # 创建提示框
          self.TipLabel = tk.Label(
              self.myTk, text="", fg="red", font=("黑体", 14), justify="center")
          self.TipLabel.grid(row=7, column=0, columnspan=5, ipadx=15)
          print("UI构建完成！")
          # 初始化持久化存储
          self.__initTobeFiledItem()
          self.myTk.mainloop()


     def __handleDate(self,isALL):  # 处理全部输出的数据
        print("开始处理数据")
        global DATAFRAME_SAVE
        global DOC_SAVE_PATH
        global DATAFRAME_SAVE
        # 1. 创建doc对象
        document = doc.Document()
        # 1.1 创建指定日期的doc对象，预留
        SpecifiedDateDoc = doc.Document()
        # 1.2 创建用户指定的日期字符串
        UserChoseDateToString =""
        # iterrows()是一个迭代器，它用于遍历DataFrame的行
        # iterrows()会返回每一行的索引和内容，其中索引是行的名称，内容是Series类型，表示一行数据
        # 2. 遍历修改模板数据
        for index, content in DATAFRAME_SAVE.iterrows():
            # 获取当前行的索引号
            all_line = len(DATAFRAME_SAVE)-1
            current_index = index
            UserChoseDate = self.SelectDateBox.get()
                
            # 判断条件：当前行的时间不为 1899-12-30 00:00:00；只有时间为空的数据在经过数据处理后会变成 1899-12-30 00:00:00
            if not str(content['时间']) == "1899-12-30 00:00:00":
                # print(type(content["时间"]))

                tempStr = self.Template
                # 取出日期左侧的0,获取用于填充至doc文件中的日期的变量
                dayStr = content.iloc[0].strftime("%m").lstrip(
                    '0')+"月"+content.iloc[0].strftime("%d").lstrip('0')+"日"
              
                if(not isALL and UserChoseDate==content.iloc[0].strftime("%d").lstrip('0')):#如果指定了日期开启并且查询到了相应的日期
                    # 对用户选择的日期赋值给顶层变量
                    UserChoseDateToString = dayStr
                else:
                    pass

                # 修改时间以及各类岗位
                tempStr = tempStr.replace("<日期>", dayStr)
                tempStr = tempStr.replace("<带班领导>", str(content["带班领导"]))
                tempStr = tempStr.replace("<总指挥>", str(content["总指挥"]))
                tempStr = tempStr.replace(
                    "<灭火救援指挥专班指挥长>", str(content["灭火救援指挥专班（备勤）"]))
                tempStr = tempStr.replace(
                    "<灭火救援指挥专班指挥助理>", str(content.iloc[12]))
                tempStr = tempStr.replace(
                    "<灭火救援指挥专班信息助理>", str(content.iloc[13]))
                tempStr = tempStr.replace(
                    "<灭火救援指挥专班通信助理>", str(content.iloc[14]))
                tempStr = tempStr.replace("<通信分队>", str(content.iloc[15]))

                # 判断是否为最后一行数据
                if not current_index == all_line:  # 不是最后一行数据
                    for position in POSISION_TITLE:
                        # 填充需要备勤信息的岗位的当天值班人员(这里填充的不是备勤人员)             
                        tempStr = tempStr.replace(f"<{position}>", str(content[position]))

                        # 填充备勤信息，如果值班表中有两个人，就执行下面这段代码
                        if "、" in str(DATAFRAME_SAVE.loc[current_index+1, position]):
                            # print("看这里："+str(DATAFRAME_SAVE.loc[current_index+1,position])+"索引为："+str(DATAFRAME_SAVE.loc[current_index+1,position].find("、")))
                            # 如果备勤的人员有两个，对原字符串进行分割，在原字符串中插入（备勤）
                            sourceStr = str( DATAFRAME_SAVE.loc[current_index+1, position])
                            split_index = str(DATAFRAME_SAVE.loc[current_index+1, position]).find("、")
                            #拼接最终的字符串
                            sourceStr = sourceStr[:split_index] +"（备勤）"+sourceStr[split_index:]
                            tempStr = tempStr.replace(f"<{position}备勤>", "、"+sourceStr)
                        else:  # 如果备勤人员有一个，就不需要进行处理，直接替换模板中的预留位即可,指挥长的备勤只可能有一个，所以在这添加对指挥长的判断
                            #使用try-except来处理异常，因为肯定会有current_index+7超出全部记录长度的情况
                            try:
                                # 若岗位为指挥长
                                if position == "指挥长":
                                    #将备勤信息替换为7天后的指挥长
                                    tempStr = tempStr.replace(f"<{position}备勤>", "、"+str(DATAFRAME_SAVE.loc[current_index+7,"指挥长"]))
                                else:
                                    #如果不是指挥长就正常替换为下一天
                                    #这里不需要判断是否为最后一行，在上面已经判断过了
                                    tempStr = tempStr.replace(f"<{position}备勤>", "、"+str(DATAFRAME_SAVE.loc[current_index+1, position]))
                            except:#当current_index+7找不到时，继续判断剩下的日期还有没有其他指挥长
                                # 获取当前的指挥长名字
                                current_zhz = str(DATAFRAME_SAVE.loc[current_index,"指挥长"])
                                # 获取在本条记录及的所有的DataFram的指挥长的值，通过tolist()转换为list
                                AfterDate = DATAFRAME_SAVE.iloc[current_index:]["指挥长"].values.tolist()
                                #使用set集合去重
                                AfterDateSet = set(AfterDate)
                                #如果指挥长存在于去重过的集合中
                                if current_zhz in AfterDateSet:
                                    #从集合内删除当前的指挥长
                                    AfterDateSet.discard(current_zhz)
                                    #再次判断set的长度，即看是不是还有其他指挥长
                                    if len(AfterDateSet)>0:
                                        #还有，就将set转换为数组，并取第一个值，也只有一个值
                                        tempStr = tempStr.replace(f"<{position}备勤>", "、"+list(AfterDateSet)[0])
                                    else:#后续没有备勤指挥长，就全部替换掉
                                        tempStr = tempStr.replace(f"<{position}备勤>（备勤）", "")

                                    
                else:  # 是最后一行数据
                      for position in POSISION_TITLE:
                        # 填充当天职位
                        tempStr = tempStr.replace(
                            f"<{position}>", str(content[position]))
                        tempStr = tempStr.replace(
                            f"<{position}备勤>", "")  # 去掉模板中的备勤
                        tempStr = tempStr.replace(f"（备勤）", "")

                # 向word中插入内容
                if(not isALL and UserChoseDate==content.iloc[0].strftime("%d").lstrip('0')):#如果是指定日期则执行下列代码
                    run = SpecifiedDateDoc.add_paragraph().add_run(tempStr)
                    #设置字体为仿宋
                    run.font.name ="仿宋"
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                else:#反之则是输出全部的
                    run = document.add_paragraph().add_run(tempStr)
                    #设置字体为仿宋
                    run.font.name ="仿宋"
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    # 插入分页符
                    document.add_page_break()
            else:  # 若时间为空，则不进行替换，跳过
                continue

        # 指定生成文件路径以及文件名,保存为word文档
        if(isALL):#全月输出 
            document.save(DOC_SAVE_PATH+"\\00.全月值班短信自动生成.docx")
            print("保存为word文档成功！")
            self.TipLabel.config(text="已生成在桌面的\n「值班短信自动生成_」文件夹内")
        else:#如果是指定日出输出，保存在下列目录中   
            SpecifiedDateDoc.save(DOC_SAVE_PATH+f"\\{UserChoseDateToString}值班信息自动生成.docx")
            print(f"{UserChoseDateToString}值班信息自动生成.docx已生成！")
            self.TipLabel.config(text=f"{UserChoseDateToString}值班信息自动生成.docx\n已创建")

     def handleFilePath(self):  # 处理excel文件，将地址赋值给全局变量
         global EXCEL_PATH
         # 弹出文件选择框
         self.TipLabel.config(text="")
         self.filePath = filedialog.askopenfilename(
             title="选择值班表文件", filetypes=[("工作簿文件", "*.xlsx")])
         # 如果用户选择了文件
         if self.filePath:
            # 1.将获取到的文件路径赋值给全局变量
            EXCEL_PATH = self.filePath
            print(f"已选择文件:{EXCEL_PATH}")
            # 调用初始化EXCEL表的函数
            if(self.__initExcel()):
                # 如果excel表初始化完成，则更改原始模板
                self.__initTemplate()
            else:
                print("------->EXCEL文件读取失败！")

         else:
            pass

     def __initExcel(self):  # 初始化excel对象
         global DATAFRAME_SAVE

         if not EXCEL_PATH == "":  # 有路径
            # 初始化dataFram对象和doc对象
           
            print("已在桌面生成「值班短信自动生成_」文件夹")

            # 1.读取原始excel文件
            data = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, header=3)

            # 2. 处理时间数据，将NaN填充为0
            data = data.fillna(0)
            #print(data["时间"])
            #print(type(data["时间"]))

            # 3. 去除名字中的空格
            for col in data.columns:
                try:
                  data[col] = data[col].str.replace(" ", "")
                except:
                    print("忽略字符串异常")

            # 4. 将时间列转换为int类型格式,方便下面转为datetime对象
            data["时间"] = data["时间"].astype(int)

            # 5. 用于将字符串、数字或其他日期时间格式的数据转换为 Pandas 中的 datetime 对象
            # unit 这个参数用于指定输入数据的时间单位 D为天，s为秒
            # origin 参数用于指定原始日期，第一个参数则是偏移量，为什么要指定这个东西，是因为原始的值班表读取出来的数据是excel的啥历史性遗留问题
            # 完成这一步后，DATAFRAME_SAVE['时间']中存储的就是 datetime对象，格式均为标准格式 xxxx-xx-xx  （年-月-日）
            data['时间'] = pd.to_datetime(data["时间"], origin="1899-12-30", unit="D")
                #print(data.columns)
                #print(data['时间'])

            print("读取excel正常！")
            # 6.将读取到的dataFram对象赋值给全局变量
            DATAFRAME_SAVE = data
            # 7.返回true
            return True
         else:
             self.TipLabel.config(text="请选择格式为 .xlsx 的文件！")
             return False
                
        
     
     def __initTemplate(self):#初始化原始模板
            self.Template = self.Template.replace("<值班专家>",self.SpecialistInput.get())
            self.Template = self.Template.replace("<值班驾驶员>",self.DriverInput.get())
            self.Template = self.Template.replace("<鲁中协作区>",self.LZInput.get())
            self.Template = self.Template.replace("<鲁南协作区>",self.LNInput.get())
            self.Template = self.Template.replace("<鲁东协作区>",self.LDInput.get())
     def WriteSpecifiedDateToDoc(self):#处理指定日期输出
         #创建桌面文件夹
         self.TipLabel.config(text="")
         if EXCEL_PATH == "" :
             self.TipLabel.config(text="请选择格式为 .xlsx 的文件！")
         else:
            # 创建桌面文件夹
            os.makedirs(DOC_SAVE_PATH, exist_ok=True)
            # 获取用户选择/输入的日期
            CaptureDate = int(self.SelectDateBox.get())
            # 如果输入的日期小于1或者大于31，就提示
            if CaptureDate < 1 or CaptureDate > 31:
                print("-------->指定日期为非法输入")
                self.TipLabel.config(text="请选择/输入1-31之间的数字")

            else:
                self.TipLabel.config(text="")
                #调用处理函数
                self.__handleDate(False)
                self.__uploadToJson()#更新持久化存储
          

     def WriteAllfiedDateToDoc(self):#写入所有的值班表
         self.TipLabel.config(text="")
         if EXCEL_PATH == "" :
             self.TipLabel.config(text="请选择格式为 .xlsx 的文件！")
         else:
            #创建桌面文件夹
            os.makedirs(DOC_SAVE_PATH, exist_ok=True)
            self.__handleDate(True)
            self.__uploadToJson()#更新持久化存储
        
     def __initTobeFiledItem(self): #初始化值守专家、值班驾驶员、鲁中协作区、鲁南协作区、鲁东协作区这些，实现本地化存储
         if os.path.exists(CREATE_FOLDER_PATH):
             pass
         else:
            os.makedirs(CREATE_FOLDER_PATH)
            print("已于 C:\\ProgramData 下创建 AutoCreateDuty 文件夹用于持久化存储")

         if os.path.exists(CREATE_JSON_PATH):#如果json文件存在，就更新到UI上
             with open(CREATE_JSON_PATH, "r", encoding="utf-8") as f:
                 item = json.load(f)
                 self.SpecialistInput.insert(0,item["值班专家"])
                 self.DriverInput.insert(0,item["值班驾驶员"])
                 self.LZInput.insert(0,item["鲁中协作区"])
                 self.LNInput.insert(0,item["鲁南协作区"])
                 self.LDInput.insert(0,item["鲁东协作区"])
         else:#如果json文件不存在，就创建一个JSON文件，并填充；注意，这个函数要在UI初始化完成后才能执行
            item = {
                "值班专家": self.SpecialistInput.get(),
                "值班驾驶员": self.DriverInput.get(),
                "鲁中协作区": self.LZInput.get(),
                "鲁南协作区": self.LNInput.get(),
                "鲁东协作区": self.LDInput.get()
            }
            with open(CREATE_JSON_PATH, "w", encoding="utf-8") as f:
                json.dump(item, f, ensure_ascii=False, indent=4)
                print("已创建用于持久化存储的json文件")

     def __uploadToJson(self):
         # 获取输入框的值
         item = {
             "值班专家": self.SpecialistInput.get(),
             "值班驾驶员": self.DriverInput.get(),
             "鲁中协作区": self.LZInput.get(),  
             "鲁南协作区": self.LNInput.get(),
             "鲁东协作区": self.LDInput.get()
         }
         # 写入JSON文件
         with open(CREATE_JSON_PATH,"w",encoding="utf-8") as f:
             json.dump(item,f,ensure_ascii=False,indent=4)
             print("已更新持久化存储：值班专家、值班驾驶员、鲁中协作区、鲁南协作区、鲁东协作区")

#创建实例对象，运行程序
start_create = CreateTable()
